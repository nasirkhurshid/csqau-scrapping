import docx
import re
from docx2pdf import convert

newDoc = docx.Document()
newTable = newDoc.add_table(0,2)
newTable.style = 'Table Grid'

pattern = "[0-9\s\S]+[AP]M"

doc = docx.Document("schedule.docx")
tables = doc.tables
for table in tables:
    for rows in table.rows:
        newTable.add_row()
        for cell in rows.cells:
            if "Time" in cell.text:
                newTable.rows[-1].cells[0].text = cell.text
            strings = re.findall(pattern, cell.text)
            if len(strings)==0:
                newTable.rows[-1].cells[1].text = cell.text
            for string in strings: 
                string = re.sub("AM", "", string)
                string = re.sub("12:30\s*PM", "12:30", string)
                if re.match(".*PM.*", string):
                    if string[1]==':':
                        t = str(int(string[0])+2)
                        string = "1"+t+string[1:]
                        if len(string)>10:
                            t = str(int(string[11])+2)
                            string = string[:11]+"1"+t+string[12:]
                    else:
                        t = str(int(string[8])+2)
                        string = string[:8]+"1"+t+string[9:]
                string = re.sub("PM", "", string)
                newTable.rows[-1].cells[0].text = ''
                newTable.rows[-1].cells[0].paragraphs[0].add_run(string).bold = True

newDoc.save("new-schedule.docx")
convert("new-schedule.docx")
print("Data written to new file")